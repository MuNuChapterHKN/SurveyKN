# -*- coding: utf-8 -*-

import re
import sys
import yaml
import pandas as pd
from shutil import copy
from pathlib import Path
from os.path import realpath
from datetime import date
from visuals import draw_pie, draw_stacked_bar
from question import add_question_to_store
from docx import Document

from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE


"""
This is the script that the user should run to generate the report templates.

Before running generate.py the necessary setup and configuration should be
completed:
    
    - the data root should be operational (setup.py)
    
    - the configuration and doctree files should be in accordance with the data
    and the expectations of the user (config.yml, doctree.yml)
    
    - [ideally] the question store should be populated beforehand (question.py)
    



"""

# The relative path to data-root-config.yml from this script
DATA_ROOT_CONFIG_REL = Path(realpath(__file__)).parent / "data-root-config.yml"
# The relative path to doctree.yml from this script
DOCTREE_REL = Path('Current Configuration/doctree.yml')
# The relative path to config.yml from this script
CONFIG_REL = Path('Current Configuration/config.yml')
# The relative path to question-store.yml from the data root
QS_FROM_ROOT = Path('AppData/question-store.yml')

def recursive_doctree_generate(dictionary, doc, working_survey, survey_id, area, config, question_store, styles_dictionary, data_root, level):
    
    """
    Gets called from generate_docs for each area in the config file. Takes the document draft and puts in
    the appropriate sections, subsections, subsubsections, the data visualizations and the comments by taking the doctree
    as a reference. Returns the document for further processing.

    Parameters
    ----------
    dictionary : Dictionary
        The dictionary that takes the recursion one level deeper, if the current dictionary has
        Questions as a value recursion stops.
    doc : docx.Document object
        The Document object that the program is currently working on
    working_survey : Pandas.DataFrame object
        A modified survey that has the question ids instead of the question texts
        for those questions that are in the working question store
    survey_id : String
        A unique identifier of the survey, generated as 'year-month'
    area : String
        The area of HKN whose report is currently being generated, the function
        generate_docs goes through the config file to determine the current areas,
        if you wish to add/remove areas, see the config file
    config : Dictionary
        The contents of the config.yml file as a Python dictionary
    question_store : Dictionary
        The question store is a dictionary whose primary purpose is to store the
        mapping { question_id : question_text } along with data relating to the
        questions' history and other internal data
        
        The question store is initially read from <data_root>/AppData/question_store.yml
    styles_dictionary : Dictionary
        The dictionary that generate_docs uses to pass the docx styles for the headings
        and the paragraphs
    data_root : pathlib.Path object
        The OS agnostic path to the data root
    level : int
        Recursion level

    Returns
    -------
    doc : docx.Document object
        The Document object that the program is currently working on

    """
    
    if level > 3:
        print("Warning: The doctree is too deep, the style of the document may be inconsistent...")
        heading_style = "Heading " + level
    else:
        heading_level = "heading" + str(level)
        heading_style = styles_dictionary[ heading_level ]
    
    for key,value in dictionary.items():
        if not isinstance(value, dict):
            # This shouldn't normally happen, this is here just in case
            doc.add_paragraph(value)
        elif "Questions" in value:
            # Add the heading
            doc.add_paragraph(key, style=heading_style)
            for question_id in value["Questions"]:
                if question_id in working_survey.columns:
                    question = question_store[question_id]["current"]
                    
                    paragraph = doc.add_paragraph("<Write here for the chart(s) below>")
                    paragraph.style = styles_dictionary["paragraph"]
                    
                    pie_chart_style = config["Pie Chart Style"]
                    
                    # Prepares the parameter for textinfo, for further details see
                    # plotly.graph_objects.Figure.update_traces
                    textinfo_value = ""
                    textinfo_percent = ""
                    textinfo_label = ""
                    
                    if pie_chart_style["On Each Slice"]["Show Value"]:
                        textinfo_value = "value+"
                    if pie_chart_style["On Each Slice"]["Show Percent"]:
                        textinfo_percent = "percent"
                    if pie_chart_style["On Each Slice"]["Show Label"]:
                        textinfo_label = "+label"
                    
                    textinfo = textinfo_value + textinfo_percent + textinfo_label
                    
                    # DEFAULT PIE CHART
                    default_pie_values = working_survey[working_survey.AREA == area][question_id].value_counts()
                    
                    default_pie_figure = draw_pie(question=question,
                                                  values=default_pie_values,
                                                  config=config,
                                                  date=survey_id,
                                                  area=area,
                                                  textinfo=textinfo)
                    
                    default_pie_path = data_root / "Visuals" / survey_id / area / (question_id + "-D.png")
                    default_pie_figure.write_image(default_pie_path.absolute().resolve().__str__(), scale=config["Document Data Visual Resolution Multiplier"])
                    doc.add_picture(default_pie_path.absolute().resolve().__str__(), width = Cm(15.0))
                    
                    
                    # CUSTOM CHARTS
                    if question_id not in config["Custom Charts Exceptions"]:
                        custom_pie_charts = {}
                        
                        if "Global Pie" in config["Custom Charts Draw"]:
                            # Global Pie Chart Variables
                            global_pie_values = working_survey[question_id].value_counts()
                            custom_pie_charts.update( { "Global Pie" : { "values" : global_pie_values,
                                                                         "date" : survey_id,
                                                                         "area" : "Associtazione" } } )
                        
                        if "Past Survey Pie" in config["Custom Charts Draw"]:
                            # Past Survey Pie Chart Variables
                            past_survey_id = config["Custom Charts"]["Past Survey Pie"]["Parameters"]
                            source_csv = data_root / "Surveys" / (past_survey_id + ".csv")
                            past_survey = pd.read_csv(source_csv)
                            past_pie_values = past_survey[past_survey.AREA == area][question_id].value_counts()
                            custom_pie_charts.update( { "Past Survey Pie" : { "values" : past_pie_values,
                                                                              "date" : past_survey_id,
                                                                              "area" : area } } )
                        
                        if "Past Survey Bar" in config["Custom Charts Draw"]:
                            # Past Survey Bar Chart Variables
                            past_survey_id_list = config["Custom Charts"]["Past Survey Bar"]["Parameters"]
                            past_bar_values = []
                            for id_code in past_survey_id_list:
                                source_csv = data_root / "Surveys" / (id_code + ".csv")
                                df = pd.read_csv(source_csv)
                                past_bar_values.append(df[df.AREA == area][question_id].value_counts())
                        
                        for chart in config["Custom Charts Draw"]:
                            if chart.strip().lower().endswith("pie"):
                                figure = draw_pie(question=question,
                                                  values=custom_pie_charts[chart]["values"],
                                                  config=config,
                                                  date=custom_pie_charts[chart]["date"],
                                                  area=custom_pie_charts[chart]["area"],
                                                  textinfo=textinfo)
                
                                chart_id = config["Custom Charts"][chart]["ID"].strip()
                                path = data_root / "Visuals" / survey_id / area / (question_id + "-" + chart_id + ".png")
                                figure.write_image(path.absolute().resolve().__str__(), scale=config["Document Data Visual Resolution Multiplier"])
                                doc.add_picture(path.absolute().resolve().__str__(), width = Cm(15.0))
                            
                            elif chart.strip().lower().endswith("bar"):
                                figure = draw_stacked_bar(question=question,
                                                          values=past_bar_values,
                                                          config=config,
                                                          dates=past_survey_id_list)
                                
                                chart_id = config["Custom Charts"][chart]["ID"].strip()
                                path = data_root / "Visuals" / survey_id / area / (question_id + "-" + chart_id + ".png")
                                figure.write_image(path.absolute().resolve().__str__(), scale=config["Document Data Visual Resolution Multiplier"])
                                doc.add_picture(path.absolute().resolve().__str__(), width = Cm(15.0))
                        
            if "Comments" in value:
                heading_level = "heading" + str(level+1)
                doc.add_paragraph("Comments", style=styles_dictionary[ heading_level ])
        
                comments = []
                for field_raw in value["Comments"]:
                    field = field_raw.strip()
                    if field not in working_survey.columns:
                        print(field + " is not in the columns of the survey, please fix the doctree...")
                        sys.exit(1)
                    comments.extend(working_survey[working_survey.AREA == area][field])
                for comment in comments:
                    if isinstance(comment, str) and len(comment) > 2:
                        doc.add_paragraph(comment, style="ListBullet")
                            
        else:
            doc.add_paragraph(key, style=styles_dictionary[ heading_level ])
            
            if key == "Comments":
                comments = []
                for field_raw in dictionary["Comments"]:
                    field = field_raw.strip()
                    if field not in working_survey.columns:
                        print(field + " is not in the columns of the survey, please fix the doctree...")
                        sys.exit(1)
                    comments.extend(working_survey[working_survey.AREA == area][field])
                for comment in comments:
                    if isinstance(comment, str) and len(comment) > 2:
                        doc.add_paragraph(comment, style="ListBullet")
            else:
                doc = recursive_doctree_generate(dictionary[key], doc, working_survey, survey_id, area, config, question_store, styles_dictionary, data_root, level+1)
            

    return doc

def generate_docs(data_root, survey_id, working_survey, config, working_doctree, question_store):
    
    """
    Goes through the config file and implements the preferences of the user.
    Creates the documents for the areas, takes care of the styles.
    Calls recursive_doctree_generate to parse the doctree and generate the visuals.

    Parameters
    ----------
    data_root : pathlib.Path object
        The OS agnostic path to the data root
    survey_id : String
        A unique identifier of the survey, generated as 'year-month'
    working_survey : Pandas.DataFrame object
        A modified survey that has the question ids instead of the question texts
        for those questions that are in the working question store
    config : Dictionary
        The contents of the config.yml file as a Python dictionary
    working_doctree : Dictionary
        This variable holds the updated and modified doctree. The questions appearing in the original
        doctree that were chosen not to be registered into the store will NOT be in the working_doctree.
        The question texts of those questions that were in the question store or were registered by the program
        in real time are replaced with their corresponding question ids
    question_store : Dictionary
        The question store is a dictionary whose primary purpose is to store the
        mapping { question_id : question_text } along with data relating to the
        questions' history and other internal data
        
        The question store is initially read from <data_root>/AppData/question_store.yml

    Returns
    -------
    None.

    """
    areas_list = config["Areas"]
    
    if config["Document Date"]["Today"]:
        date_string = date.today().strftime("%B %Y")
    else:
        months = { "01" : "January", "02" : "February", "03" : "March", "04" : "April",
                  "05" : "May", "06" : "June", "07" : "July", "08" : "August", "09" : "September",
                  "10" : "October", "11" : "November", "12" : "December" }
        date_string = months[survey_id.split("-")[1]] + " " + survey_id.split("-")[0]
        
    print("\nThis may take a minute, sit back and relax...\n")

    
    for area in areas_list:
        print("Generating report for Area " + area + "...")
        doc = Document()
        styles = doc.styles
        
        # Title Style
        config_document_title = config["Document Title"]
        title_style = styles.add_style("title style", WD_STYLE_TYPE.PARAGRAPH)
        title_style.base_style = styles["Heading 1"]
        title_style.font.name = config_document_title["Font"]
        title_style.font.size = Pt(config_document_title["Font Size"])
        title_style.font.bold = config_document_title["Bold"]
        color_hex = config_document_title["Font Color"].strip("#")
        rgb_tuple = tuple(int(color_hex[i:i+2], 16) for i in (0, 2, 4))
        title_style.font.color.rgb = RGBColor(*rgb_tuple)
        
        # Date Style
        config_document_date = config["Document Date"]
        date_style = styles.add_style("date style", WD_STYLE_TYPE.PARAGRAPH)
        date_style.font.name = config_document_date["Font"]
        date_style.font.size = Pt(config_document_date["Font Size"])
        date_style.font.bold = config_document_date["Bold"]
        color_hex = config_document_date["Font Color"].strip("#")
        rgb_tuple = tuple(int(color_hex[i:i+2], 16) for i in (0, 2, 4))
        date_style.font.color.rgb = RGBColor(*rgb_tuple)
        
        # Metadata Style
        config_document_metadata = config["Document Metadata"]
        metadata_style = styles.add_style("metadata style", WD_STYLE_TYPE.PARAGRAPH)
        metadata_style.font.name = config_document_metadata["Font"]
        metadata_style.font.size = Pt(config_document_metadata["Font Size"])
        metadata_style.font.bold = config_document_metadata["Bold"]
        color_hex = config_document_metadata["Font Color"].strip("#")
        rgb_tuple = tuple(int(color_hex[i:i+2], 16) for i in (0, 2, 4))
        metadata_style.font.color.rgb = RGBColor(*rgb_tuple)
        
        # Disclaimer Style
        config_document_disclaimer = config["Document Disclaimer"]
        disclaimer_style = styles.add_style("disclaimer style", WD_STYLE_TYPE.PARAGRAPH)
        disclaimer_style.font.name = config_document_disclaimer["Font"]
        disclaimer_style.font.size = Pt(config_document_disclaimer["Font Size"])
        disclaimer_style.font.bold = config_document_disclaimer["Bold"]
        color_hex = config_document_disclaimer["Font Color"].strip("#")
        rgb_tuple = tuple(int(color_hex[i:i+2], 16) for i in (0, 2, 4))
        disclaimer_style.font.color.rgb = RGBColor(*rgb_tuple)
        
        # Heading1 Style
        config_document_heading1 = config["Document Section Heading"]
        heading1_style = styles.add_style("heading1 style", WD_STYLE_TYPE.PARAGRAPH)
        heading1_style.base_style = styles["Heading 1"]
        heading1_style.font.name = config_document_heading1["Font"]
        heading1_style.font.size = Pt(config_document_heading1["Font Size"])
        heading1_style.font.bold = config_document_heading1["Bold"]
        color_hex = config_document_heading1["Font Color"].strip("#")
        rgb_tuple = tuple(int(color_hex[i:i+2], 16) for i in (0, 2, 4))
        heading1_style.font.color.rgb = RGBColor(*rgb_tuple)
        
        # Heading2 Style
        config_document_heading2 = config["Document Subsection Heading"]
        heading2_style = styles.add_style("heading2 style", WD_STYLE_TYPE.PARAGRAPH)
        heading2_style.base_style = styles["Heading 2"]
        heading2_style.font.name = config_document_heading2["Font"]
        heading2_style.font.size = Pt(config_document_heading2["Font Size"])
        heading2_style.font.bold = config_document_heading2["Bold"]
        color_hex = config_document_heading2["Font Color"].strip("#")
        rgb_tuple = tuple(int(color_hex[i:i+2], 16) for i in (0, 2, 4))
        heading2_style.font.color.rgb = RGBColor(*rgb_tuple)
        
        # Heading3 Style
        config_document_heading3 = config["Document Subsection Heading"]
        heading3_style = styles.add_style("heading3 style", WD_STYLE_TYPE.PARAGRAPH)
        heading3_style.base_style = styles["Heading 3"]
        heading3_style.font.name = config_document_heading3["Font"]
        heading3_style.font.size = Pt(config_document_heading3["Font Size"])
        heading3_style.font.bold = config_document_heading3["Bold"]
        color_hex = config_document_heading3["Font Color"].strip("#")
        rgb_tuple = tuple(int(color_hex[i:i+2], 16) for i in (0, 2, 4))
        heading3_style.font.color.rgb = RGBColor(*rgb_tuple)
        
        # Paragraph Style
        config_document_paragraph = config["Document Paragraph"]
        paragraph_style = styles.add_style("paragraph style", WD_STYLE_TYPE.PARAGRAPH)
        paragraph_style.font.name = config_document_paragraph["Font"]
        paragraph_style.font.size = Pt(config_document_paragraph["Font Size"])
        paragraph_style.font.bold = config_document_paragraph["Bold"]
        color_hex = config_document_paragraph["Font Color"].strip("#")
        rgb_tuple = tuple(int(color_hex[i:i+2], 16) for i in (0, 2, 4))
        paragraph_style.font.color.rgb = RGBColor(*rgb_tuple)
        
        
        document_title = doc.add_heading(config_document_title["Text"], level=1)
        document_title.style = title_style
        if config_document_title["Middle Aligned"]:
            document_title.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        document_date = doc.add_paragraph(date_string)
        document_date.style = date_style
        
        survey_area = working_survey[working_survey.AREA == area]
        metadata_string = "Relatore: \n" "Area: " + area + "\nCompilazioni ottenute: " + str(len(survey_area.index))
        document_metadata = doc.add_paragraph(metadata_string)
        document_metadata.style = metadata_style
        
        document_disclaimer = doc.add_paragraph(config["Document Disclaimer"]["Text"])
        document_disclaimer.style = disclaimer_style
        
        styles_dictionary = { "heading1" : heading1_style,
                              "heading2" : heading2_style,
                              "heading3" : heading3_style,
                              "paragraph" : paragraph_style }
        doc = recursive_doctree_generate(working_doctree, doc, working_survey, survey_id, area, config, question_store, styles_dictionary, data_root, 1)
        
        if config["Document Conclusion Tree"]:
            document_conclusion = doc.add_heading("Conclusion", level=1)
            document_conclusion.style = heading1_style
            conclusion_headings = []
            conclusion_paragraphs = []
            for section in working_doctree:
                conclusion_headings.append(doc.add_heading(section, level=2))
                conclusion_paragraphs.append(doc.add_paragraph(section + " summary"))
            
            for heading in conclusion_headings:
                heading.style = heading2_style
            
            
        doc.save(data_root / "Templates" / survey_id / (area + ".docx"))
        
    return

def process_survey(survey_id, question_store, working_inverted_question_store, survey):
    
    """
    First strips the questions in the survey from leading, trailing white space and
    any question numbers that may be present (i.e '2) ' ).
    Then, for every column name in the DataFrame, it checks to see if the question
    is in the working_inverted_question_store, if the question is there, the question
    text is replaced by the question id to facilitate further processing and future
    reference.

    Parameters
    ----------
    survey_id : String
        A unique identifier of the survey, generated as 'year-month'
    question_store : Dictionary
        The question store is a dictionary whose primary purpose is to store the
        mapping { question_id : question_text } along with data relating to the
        questions' history and other internal data
        
        The question store is initially read from <data_root>/AppData/question_store.yml
    working_inverted_question_store : Dictionary
        The working_inverted_question_store contains the same type of mapping that
        inverted_question_store has, however the working_inverted_question_store may also contain
        newly registered questions that were taken from the doctree. Also any question not found
        in the doctree will NOT be in the working_inverted_question_store even if it is in
        the inverted_question_store.
    survey : Pandas.DataFrame object
        The contents of the csv file as a Pandas Data Frame object

    Returns
    -------
    question_store : Dictionary
        The question store is a dictionary whose primary purpose is to store the
        mapping { question_id : question_text } along with data relating to the
        questions' history and other internal data
        
        The question store is initially read from <data_root>/AppData/question_store.yml
    working_survey : Pandas.DataFrame object
        A modified survey that has the question ids instead of the question texts
        for those questions that are in the working question store

    """
    
    new_column_names = []
    for column in survey.columns:
        # Remove the leading 'n) ' for the questions that have it
        question = column.strip('1234567890) ')
        
        # If this is a question that's in the working question store updates
        # the question's history in the store and replaces the question
        # text in the csv file with its question id
        if question in working_inverted_question_store:
            question_id = working_inverted_question_store[question]
            question_store[question_id].update( { survey_id : question } )
            question = question_id
        
        new_column_names.append(question)
   
    survey.columns = new_column_names
    working_survey = survey
    
    return question_store, working_survey


def recursive_doctree_question_store(dictionary, question_store, inverted_question_store, working_inverted_question_store, survey_id):
    
    """
    Recursively visits the doctree, checks if the questions are in the question store, if a question is not in the question store
    recursive_doctree_question_store prompts the user and offers to register the question into the question store, otherwise the
    question is ignored. For those questions that were in the question store or that were registered by the function, the question
    texts are replaced by the corresponding question ids to create the working_doctree. During this process the function also builds
    the working_inverted_question_store.

    Parameters
    ----------
    dictionary : Dictionary
        The dictionary that takes the recursion one level deeper, if the current dictionary has
        Questions as a value recursion stops.
    question_store : Dictionary
        The question store is a dictionary whose primary purpose is to store the
        mapping { question_id : question_text } along with data relating to the
        questions' history and other internal data
        
        The question store is initially read from <data_root>/AppData/question_store.yml
    inverted_question_store : Dictionary
        The inverted question store holds the inverted mapping { question_text : question_id }
        to facilitate processing
    working_inverted_question_store : Dictionary
        The working_inverted_question_store contains the same type of mapping that
        inverted_question_store has, however the working_inverted_question_store may also contain
        newly registered questions that were taken from the doctree. Also any question not found
        in the doctree will NOT be in the working_inverted_question_store even if it is in
        the inverted_question_store. Originally passed as an empty dictionary from the wrapper.
    survey_id : String
        A unique identifier of the survey, generated as 'year-month'

    Returns
    -------
    question_store : Dictionary
        The question store is a dictionary whose primary purpose is to store the
        mapping { question_id : question_text } along with data relating to the
        questions' history and other internal data
        
        The question store is initially read from <data_root>/AppData/question_store.yml
    working_inverted_question_store : Dictionary
        The working_inverted_question_store contains the same type of mapping that
        inverted_question_store has, however the working_inverted_question_store may also contain
        newly registered questions that were taken from the doctree. Also any question not found
        in the doctree will NOT be in the working_inverted_question_store even if it is in
        the inverted_question_store.
    new_dictionary : Dictionary
        This dictionary is used to build the working_doctree while going back up from recursion.
        Has the same structure as doctree, it won't contain those questions the user chose to
        ignore, it will have the question ids instead of the question texts.

    """
    
    # new_dictionary is used for building the working_doctree while going back from recursion
    new_dictionary = {}
    for key,value in dictionary.items():
        if not isinstance(value, dict):
            pass
        elif "Questions" in value:
            # question_id_list will hold the question ids that will replace the question texts
            # in the working doctree
            question_id_list = []
            question_list = value["Questions"]
            
            for question_raw in question_list:
                question = question_raw.strip()
                if question in inverted_question_store:
                    question_id_list.append(inverted_question_store[question])
                    # See docstring for the purpose of the working_inverted_question_store
                    working_inverted_question_store.update( { question : inverted_question_store[question] } )
                    
                else:
                    print("\n\n\"" + question + "\"\n Is not in the question store.\n")
                    print("\nIf this question isn't registered into the store, it will be ignored while generating the reports.\n")
                    print("Would you like to create a new entry for this question now?")
                    answer = input("[yes/no] > ")
    
                    if answer.lower() == "yes":
                        # Function call to question.add_question_to_store from question.py
                        question_store, question_id = add_question_to_store(question_store, question)
                        question_id_list.append(question_id)
                        working_inverted_question_store.update( { question : question_id } )
                    else:
                        print("\nThis question will not appear in the generated reports.\n\n\n")
                
            # The question ids replace the question texts in the working_doctree
            if "Comments" in value:
                new_value = { "Questions" :  question_id_list, "Comments" : value["Comments"]}
            else:
                new_value = { "Questions" :  question_id_list }
            
            # Appends the question ids to the new dictionary
            new_dictionary.update({ key : new_value })
            
        else:
            question_store, working_inverted_question_store, new_value = recursive_doctree_question_store(dictionary[key],
                                                                                                      question_store,
                                                                                                      inverted_question_store,
                                                                                                      working_inverted_question_store,
                                                                                                      survey_id)
            # Appends the question ids that were returned from the lower levels
            new_dictionary.update({ key : new_value })

    return question_store, working_inverted_question_store, new_dictionary

def process_doctree(survey_id, question_store, inverted_question_store, doctree):
    
    """
    This is nothing but a wrapper function for recursive_doctree_question_store

    Parameters
    ----------
    survey_id : String
        A unique identifier of the survey, generated as 'year-month'
    question_store : Dictionary
        The question store is a dictionary whose primary purpose is to store the
        mapping { question_id : question_text } along with data relating to the
        questions' history and other internal data
        
        The question store is initially read from <data_root>/AppData/question_store.yml
    inverted_question_store : Dictionary
        The inverted question store holds the inverted mapping { question_text : question_id }
        to facilitate processing
    doctree : Dictionary
        The doctree holds a tree structure that represents the titles, subtitles,
        subsubtitles,... and questions of the documents that are to be generated

    Returns
    -------
    question_store : Dictionary
        The question store is a dictionary whose primary purpose is to store the
        mapping { question_id : question_text } along with data relating to the
        questions' history and other internal data
        
        The question store is initially read from <data_root>/AppData/question_store.yml
    working_inverted_question_store : Dictionary
        The working_inverted_question_store contains the same type of mapping that
        inverted_question_store has, however the working_inverted_question_store may also contain
        newly registered questions that were taken from the doctree. Also any question not found
        in the doctree will NOT be in the working_inverted_question_store even if it is in
        the inverted_question_store.
    working_doctree : Dictionary
        This variable holds the updated and modified doctree. The questions appearing in the original
        doctree that were chosen not to be registered into the store will NOT be in the working_doctree.
        The question texts of those questions that were in the question store or were registered by the program
        in real time are replaced with their corresponding question ids

    """
    
    working_inverted_question_store = {}
    question_store, working_inverted_question_store, working_doctree = recursive_doctree_question_store(doctree,
                                                                                                        question_store,
                                                                                                        inverted_question_store,
                                                                                                        working_inverted_question_store,
                                                                                                        survey_id)
    
    return question_store, working_inverted_question_store, working_doctree

def create_directories(data_root, survey_id, config):
    
    """
    Creates the necessary directories for this survey and its data using pathlib.Path.mkdir
    
    If generate.py is called more than once on the same survey, no errors will
    be raised, the old files will be overwritten

    Parameters
    ----------
    data_root : pathlib.Path object
        The OS agnostic path to the data root
    survey_id : String
        A unique identifier of the survey, generated as 'year-month'
    config : Dictionary
        The contents of the config.yml file as a Python dictionary

    Returns
    -------
    None.

    """
    
    visuals_path = data_root / "Visuals" / survey_id
    visuals_path.mkdir(exist_ok=True)
    for area in config["Areas"]:
        path = visuals_path / area
        path.mkdir(exist_ok=True)
    
    templates_path = data_root / "Templates" / survey_id
    templates_path.mkdir(exist_ok=True) 
    return

def verify_and_register_csv(data_root):
    
    """
    Checks the validity of the given path that should lead to a .csv file
    
    If the path checks out, the user is asked to provide the survey year and month
    These values are used to assign the survey ID as 'year-month'
    
    After determining the survey ID, the program creates an exact copy of the csv file
    using shutils.copy in <data_root>/Surveys named <survey_id>-original.csv
    
    If generate.py was called with the same survey before, it overwrites the previously
    copied csv file in <data_root>/Surveys

    Parameters
    ----------
    data_root : pathlib.Path object
        The OS agnostic path to the data root

    Returns
    -------
    survey_id : String
        A unique identifier of the survey, generated as 'year-month'
    survey : Pandas.DataFrame object
        The contents of the csv file as a Pandas Data Frame object

    """
    
    source_csv = Path(sys.argv[1])
    
    if not source_csv.is_file():
        print("The given path doesn't point to a file.")
        print("Please run as generate.py <path/to/survey/results.csv>")
        sys.exit(1)
    
    if not source_csv.suffix == ".csv":
        print("The given path doesn't point to a .csv file, this program only accepts .csv files.")
        print("Please run as generate.py <path/to/survey/results.csv>")
        sys.exit(1)
    
    # Regular expression for checking the general formatting of the year
    rex = re.compile("^[0-9]{4}$")
    
    # Makeshift 'do while' loop
    while True:
        year = input("\nYear of survey: ")
        if rex.match(year):
            break
        print("\nNot a valid year, please enter the year represented as 4 digits\n\n")
    
    # Regular expression for checking the general formatting of the month
    rex = re.compile("^[0-9]{2}$")
    
    # Makeshift 'do while' loop
    while True:
        month = input("Month of survey [0-12]: ")
        if rex.match(month):
            break
        print("\nNot a valid month, please enter the month represented as 2 digits ( ie. for August enter 08 )\n\n")
    
    survey_id = year + "-" + month
    
    # We add 'original' to the end because we will save another version
    # of the survey with its plain name, it will have the question ids
    # instead of the question texts
    csv_filename = survey_id + "-original.csv"
    
    target_csv = data_root /  "Surveys" / csv_filename
    copy(source_csv, target_csv)
    
    # Load the survey as a Pandas DataFrame
    survey = pd.read_csv(source_csv)
    
    return survey_id, survey

def prompt_user_for_prerequisites():
    
    """
    Prompts the user to see if they have made the necessary changes to the
    configuration files: config.yml and doctree.yml
    
    Asks for confirmation before resuming the program

    Returns
    -------
    None.

    """
    
    print("\nThis script will generate survey response analysis templates in .docx format from the given .csv file containing the results.")
    print("The output will heavily depend on config.yml and doctree.yml that can be found in the folder named Current Configuration.\n")
    print("Have you edited these files to your liking?")
    print("Do you wish to proceed?")
    answer = input("[yes/no] > ")
    
    if answer.lower() == "yes":
        return
    
    print("\nExiting, no changes were made.\n")
    sys.exit(0)
    

def main():
    
    if len(sys.argv) != 2 or sys.argv[1] == "help":
        print(main.__doc__)
        return
    
    prompt_user_for_prerequisites()
    
    with open(DATA_ROOT_CONFIG_REL, "r") as fd:
         data_root = Path(yaml.safe_load(fd)['root'])
    
    survey_id, survey = verify_and_register_csv(data_root)
    
    with open(CONFIG_REL, "r") as fd:
         config = yaml.safe_load(fd)
    
    with open(DOCTREE_REL, "r") as fd:
         doctree = yaml.safe_load(fd)
    
    with open(data_root / QS_FROM_ROOT, "r") as fd:
         question_store = yaml.safe_load(fd)
         
    create_directories(data_root, survey_id, config)
    
    # Copy the config file and the doctree into the data root for future reference
    copy(DOCTREE_REL, data_root / "Templates" / survey_id / "doctree.yml")
    copy(CONFIG_REL, data_root / "Templates" / survey_id / "config.yml")
    
    inverted_question_store = {}
    
    # Creates a dictionary with the mapping (question_text : question_id)
    for k,v in question_store.items():
        if not k == "NEXTINLINE" and not k == "COUNT":
            inverted_question_store.update({ v["current"] : k })
    
    # Goes through the doctree and replaces the questions' text with their question ids
    question_store, working_inverted_question_store, working_doctree = process_doctree(survey_id, question_store, inverted_question_store, doctree)
    
    # Goes through the survey data and replaces those questions that show up in
    # the doctree with their corresponding question ids
    question_store, working_survey = process_survey(survey_id, question_store, working_inverted_question_store, survey)
    
    print("\n")
    
    # Generates the visuals and the templates
    generate_docs(data_root, survey_id, working_survey, config, working_doctree, question_store)
    
    print("\nCleaning up...")
    
    csv_filename = survey_id + ".csv"
    working_survey.to_csv(data_root /  "Surveys" / csv_filename)
    
    with open(data_root / QS_FROM_ROOT, "w") as fd:
            yaml.dump(question_store, fd)
    
    print("\nSuccessfully generated the report templates!\n")
    
    return

if __name__ == "__main__":
    main()
